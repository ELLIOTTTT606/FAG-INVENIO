"""Synthetic PAC and GEG GALLETTI fiches for tests.

These mimic the structure described in `docs/architecture.md` and the
spec the Product Owner shared in the Sprint 1 thread:

  - Header "PROPOSITION DE PROJET"
  - Designation paragraph ("PLP052HS2B A000CE000I00110 ...")
  - Conditions de fonctionnement (Refroidissement, Chauffage, Acoustique, UNI EN 14511)
  - Performances (Refroidissement, Chauffage)
  - Donnees generales / Common Data
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument


def _add_section_table(document: DocxDocument, title: str, rows: list[tuple[str, str, str]]) -> None:
    document.add_paragraph(title)
    table = document.add_table(rows=len(rows), cols=3)
    for i, (label, unit, value) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = unit
        cells[2].text = value


def _add_general_table(document: DocxDocument, rows: list[tuple[str, str, str]]) -> None:
    _add_section_table(document, "Donnees generales", rows)


def build_pac_document() -> DocxDocument:
    """Build a PAC fiche close to a real GALLETTI export (PLP052HS2B)."""

    document = Document()
    document.add_paragraph("PROPOSITION DE PROJET")
    document.add_paragraph("Reference: PLP052HS2B A000CE000I00110 0000000I000000000000")
    document.add_paragraph(
        "Agence : FRANCE AIR - 383 rue des Barronnieres - 01700 Beynost - "
        "Donnees : 13-02-2026"
    )

    _add_section_table(
        document,
        "Refroidissement",
        [
            ("Temp. entree eau utilisation", "C", "12.0"),
            ("Temp. sortie eau utilisation", "C", "7.0"),
            ("Glycol cote utilisation", "%", "0"),
            ("Temp. air exterieur source", "C", "35.0"),
            ("Humidite relative air source", "%", "40"),
            ("Pourcentage de charge", "%", "100"),
        ],
    )
    _add_section_table(
        document,
        "Chauffage",
        [
            ("Temp. entree eau utilisation", "C", "40.0"),
            ("Temp. sortie eau utilisation", "C", "45.0"),
            ("Glycol cote utilisation", "%", "0"),
            ("Temp. air exterieur source", "C", "7.0"),
            ("Humidite relative air source", "%", "87"),
            ("Pourcentage de charge", "%", "100"),
        ],
    )
    _add_section_table(
        document,
        "Donnees acoustiques",
        [
            ("Distance en champ libre", "m", "10.0"),
            ("Facteur de directionnalite", "", "2"),
        ],
    )
    _add_section_table(
        document,
        "Norme UNI EN 14511",
        [
            ("Calculs selon la norme UNI EN 14511", "", "Oui"),
            ("UNI EN 14511 Version", "", "UNI EN 14511 - 2022"),
        ],
    )

    _add_section_table(
        document,
        "Refroidissement (Performances)",
        [
            ("Puissance frigorifique", "kW", "41.7"),
            ("Debit eau utilisateur", "l/h", "7155"),
            ("Perte de charge utilisateur", "kPa", "30"),
            ("Puissance absorbee compresseurs", "kW", "16.0"),
            ("Courant absorbe compresseurs", "A", "25.6"),
            ("Puissance absorbee totale", "kW", "16.6"),
            ("Courant absorbe total", "A", "28.8"),
            ("EER", "W/W", "2.50"),
            ("EER UNI EN 14511", "W/W", "2.48"),
            ("SEER", "Wh/Wh", "4.15"),
        ],
    )
    _add_section_table(
        document,
        "Chauffage (Performances)",
        [
            ("Puissance calorifique", "kW", "52.3"),
            ("Debit eau utilisateur", "l/h", "9087"),
            ("Perte de charge utilisateur", "kPa", "48"),
            ("Puissance absorbee compresseurs", "kW", "14.8"),
            ("Courant absorbe compresseurs", "A", "23.8"),
            ("Puissance absorbee totale", "kW", "15.5"),
            ("Courant absorbe total", "A", "27.0"),
            ("COP", "W/W", "3.37"),
            ("COP UNI EN 14511", "W/W", "3.32"),
            ("SCOP", "Wh/Wh", "4.35"),
            ("Eta s", "", "171.0"),
            ("Classe saisonniere chauffage", "", "A++"),
        ],
    )

    _add_general_table(
        document,
        [
            ("Max courant absorbe (FLA)", "A", "56"),
            ("Courant de demarrage (LRA)", "A", "57"),
            ("Niveau de puissance acoustique Lw", "dB(A)", "83"),
            ("Niveau de pression acoustique Lp", "dB(A)", "55"),
            ("Debit d'air source", "m3/h", "16259"),
            ("Nombre de ventilateurs", "", "2"),
            ("Refrigerant", "", "R290"),
            ("GWP", "", "3"),
            ("Poids sans options", "kg", "500"),
            ("Alimentation", "", "400 / 3+N / 50"),
        ],
    )

    return document


def build_geg_document() -> DocxDocument:
    """Build a GEG fiche close to a real GALLETTI export (VLS202CS0A).

    GEG fiches have a single Refroidissement section and no Chauffage block.
    """

    document = Document()
    document.add_paragraph("PROPOSITION DE PROJET")
    document.add_paragraph("Reference: VLS202CS0A B000CE000I00220 0000000I000000000000")
    document.add_paragraph(
        "Agence : FRANCE AIR - 383 rue des Barronnieres - 01700 Beynost - "
        "Donnees : 13-02-2026"
    )

    _add_section_table(
        document,
        "Refroidissement",
        [
            ("Temp. entree eau utilisation", "C", "12.0"),
            ("Temp. sortie eau utilisation", "C", "7.0"),
            ("Glycol cote utilisation", "%", "0"),
            ("Temp. air exterieur source", "C", "35.0"),
            ("Humidite relative air source", "%", "40"),
            ("Pourcentage de charge", "%", "100"),
        ],
    )
    _add_section_table(
        document,
        "Donnees acoustiques",
        [
            ("Distance en champ libre", "m", "10.0"),
            ("Facteur de directionnalite", "", "2"),
        ],
    )
    _add_section_table(
        document,
        "Norme UNI EN 14511",
        [
            ("Calculs selon la norme UNI EN 14511", "", "Oui"),
            ("UNI EN 14511 Version", "", "UNI EN 14511 - 2022"),
        ],
    )

    _add_section_table(
        document,
        "Refroidissement (Performances)",
        [
            ("Puissance frigorifique", "kW", "210.0"),
            ("Debit eau utilisateur", "l/h", "36000"),
            ("Perte de charge utilisateur", "kPa", "45"),
            ("Puissance absorbee compresseurs", "kW", "72.0"),
            ("Courant absorbe compresseurs", "A", "110.0"),
            ("Puissance absorbee totale", "kW", "75.0"),
            ("Courant absorbe total", "A", "118.0"),
            ("EER", "W/W", "2.80"),
            ("EER UNI EN 14511", "W/W", "2.78"),
            ("SEER", "Wh/Wh", "4.10"),
        ],
    )

    _add_general_table(
        document,
        [
            ("Max courant absorbe (FLA)", "A", "180"),
            ("Courant de demarrage (LRA)", "A", "210"),
            ("Niveau de puissance acoustique Lw", "dB(A)", "88"),
            ("Niveau de pression acoustique Lp", "dB(A)", "60"),
            ("Debit d'air source", "m3/h", "52000"),
            ("Nombre de ventilateurs", "", "4"),
            ("Refrigerant", "", "R454B"),
            ("GWP", "", "466"),
            ("Poids sans options", "kg", "900"),
            ("Alimentation", "", "400 / 3+N / 50"),
        ],
    )

    return document


def write_pac(path: Path) -> Path:
    document = build_pac_document()
    document.save(str(path))
    return path


def write_geg(path: Path) -> Path:
    document = build_geg_document()
    document.save(str(path))
    return path
