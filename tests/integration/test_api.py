"""Integration test: upload a DOCX through the FastAPI endpoint."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from src.api.main import app


@pytest.fixture()
def sample_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Reference: PLP052HS2B A000CE000I00110 0000000I000000000000")
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Puissance frigorifique"
    table.rows[0].cells[1].text = "kW"
    table.rows[0].cells[2].text = "41,7"
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_health_endpoint() -> None:
    client = TestClient(app)
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_parse_docx_endpoint_returns_canonical_payload(sample_docx_bytes: bytes) -> None:
    client = TestClient(app)
    files = {
        "file": (
            "sample.docx",
            sample_docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    response = client.post("/parse/docx", files=files)
    assert response.status_code == 200, response.text
    body = response.json()
    assert body["data"]["model"] == "PLP"
    assert body["data"]["performance"]["cooling_power_kW"] == 41.7


def test_parse_docx_endpoint_rejects_wrong_extension() -> None:
    client = TestClient(app)
    files = {"file": ("bad.txt", b"hello", "text/plain")}
    response = client.post("/parse/docx", files=files)
    assert response.status_code == 415


def test_parse_real_sample_file_if_present() -> None:
    sample = Path("examples/sample_galletti.docx")
    if not sample.exists():
        pytest.skip("examples/sample_galletti.docx not provided yet")
    client = TestClient(app)
    with sample.open("rb") as f:
        response = client.post(
            "/parse/docx",
            files={"file": (sample.name, f.read(), "application/octet-stream")},
        )
    assert response.status_code == 200
